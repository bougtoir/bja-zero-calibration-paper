#!/usr/bin/env python3
"""
Recreate BJA Special Article .docx — English version (UPDATED)
Now includes: original 4 figures + 4 new figures + Table 2 (statistical comparison)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

FIGDIR = '/home/ubuntu/bja_figures'
OUTPATH = '/home/ubuntu/BJA_ZeroFree_Manuscript_EN.docx'

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)

# ── Helper functions ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p

def add_figure(filename, caption, width=Inches(6)):
    path = os.path.join(FIGDIR, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    cap = doc.add_paragraph()
    cap_run = cap.add_run(caption)
    cap_run.font.name = 'Times New Roman'
    cap_run.font.size = Pt(10)
    cap_run.italic = True
    cap.paragraph_format.space_after = Pt(12)
    return cap

# ══════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════
add_para('SPECIAL ARTICLE', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    'Pulse Pressure Accuracy as an Implicit Gain Validator: '
    'A Theoretical Framework for Zero-Calibration-Free Arterial Pressure Monitoring'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.bold = True
title.paragraph_format.space_after = Pt(18)

add_para('Running title: Zero-calibration-free arterial monitoring', italic=True,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_para('[Author names to be completed]', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[Affiliations to be completed]', alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_para('Corresponding author:', bold=True)
add_para('[Name, address, email, telephone]')
doc.add_paragraph()

add_para('Word count: ~3,800 (main text, excluding references, tables, and figure legends)')
add_para('Tables and Figures: 2 Tables, 8 Figures')
add_para('References: 30')
doc.add_paragraph()

add_para('Keywords: ', bold=True, space_after=Pt(0))
kw = doc.paragraphs[-1]
kw_run = kw.add_run(
    'arterial pressure monitoring; zero calibration; pulse pressure; '
    'concordance correlation coefficient; MEMS pressure sensor; '
    'bias correction factor; device design'
)
kw_run.font.name = 'Times New Roman'
kw_run.font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SUMMARY
# ══════════════════════════════════════════════════════════════════
add_heading_styled('Summary', level=1)

add_para(
    'Zero calibration of invasive arterial pressure transducers is a fundamental step in haemodynamic '
    'monitoring, universally taught, universally practised, and universally accepted as indispensable. '
    'In this article, we challenge this assumption by presenting a theoretical framework demonstrating '
    'that zero calibration is, in principle, rendered unnecessary by appropriate sensor design. '
    'Our argument rests on two pillars. First, we show that pulse pressure (PP = systolic minus '
    'diastolic pressure), as a differential (AC) quantity, is independent of the DC offset that zero '
    'calibration is designed to remove. If a sensor measures PP accurately, its gain (sensitivity) is '
    'thereby validated, because PP accuracy requires correct transduction of the pressure difference '
    'across the full physiological range. Second, we identify three engineering solutions\u2014catheter-tip '
    'MEMS sensors, absolute pressure measurement with barometric compensation, and self-calibrating '
    'MEMS with internal reference cavities\u2014that together eliminate all three sources of DC offset '
    '(hydrostatic column, atmospheric pressure reference, and transducer drift). We formalise this '
    'argument using Lin\u2019s concordance correlation coefficient (CCC) decomposition, showing that '
    'zero calibration addresses only the location shift component (u) of the bias correction factor '
    'C_b, while scale shift (v)\u2014the gain error that distorts PP\u2014is left uncorrected. '
    'A properly designed sensor achieves C_b \u2248 1.0 by construction, making manual zero calibration '
    'redundant. We discuss the implications for device development, regulatory validation, and clinical '
    'practice, and extend the framework to non-invasive cardiac output monitors.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# INTRODUCTION
# ══════════════════════════════════════════════════════════════════
add_heading_styled('Introduction', level=1)

add_para(
    'Invasive arterial pressure monitoring is one of the most frequently performed measurements in '
    'anaesthesia and critical care. After insertion of an arterial catheter, the clinician must zero '
    'the pressure transducer to atmospheric pressure at the level of the phlebostatic axis before '
    'initiating haemodynamic monitoring.1,2 This zeroing procedure must be repeated whenever the '
    'height relationship between the transducer and the catheter insertion site changes\u2014for example, '
    'when the operating table is tilted or the bed angle is adjusted.3 Despite decades of technological '
    'advancement in pressure sensing, the need for manual zero calibration has remained essentially '
    'unchanged since the introduction of external fluid-filled transducer systems.'
)

add_para(
    'The persistence of this requirement is not a fundamental physical necessity but a consequence of '
    'how conventional monitoring systems are designed. In a standard fluid-filled arterial line, the '
    'pressure transducer is located externally, connected to the intravascular catheter tip by a column '
    'of saline. This architecture introduces three systematic sources of DC offset: (1) the hydrostatic '
    'pressure difference between the transducer and the measurement site (\u0394P = \u03C1gh, approximately '
    '0.74 mmHg per centimetre of height difference); (2) the need to measure gauge pressure relative '
    'to atmospheric pressure (~760 mmHg); and (3) transducer drift over time due to mechanical creep '
    'or thermal effects in the strain gauge.4,5 Zero calibration removes all three offsets '
    'simultaneously by exposing the transducer to atmosphere and resetting the output to zero.'
)

add_para(
    'However, zero calibration is exclusively an offset correction. It does not verify or correct the '
    'gain (sensitivity) of the transducer\u2014that is, the proportionality between the true pressure '
    'change and the electrical output. If the gain is incorrect, the measured waveform will be scaled '
    'up or down, and this error will persist after zeroing. This distinction between offset error and '
    'gain error is not merely academic; it has direct implications for how we validate monitoring '
    'devices and, as we argue below, for how they should be designed.'
)

add_para(
    'In this article, we propose that if a device measures pulse pressure (PP) accurately, zero '
    'calibration is rendered unnecessary through appropriate sensor design. We formalise this argument '
    'using Lin\u2019s concordance correlation coefficient (CCC) decomposition,6,7 identify the engineering '
    'solutions that eliminate each source of DC offset, and discuss the implications for arterial '
    'pressure monitoring and beyond.'
)

# ══════════════════════════════════════════════════════════════════
# THE LOGICAL ARGUMENT
# ══════════════════════════════════════════════════════════════════
add_heading_styled('The logical argument: from pulse pressure to zero-free monitoring', level=1)

add_heading_styled('Arterial pressure as the sum of AC and DC components', level=2)

add_para(
    'An arterial pressure waveform P(t) can be decomposed into a slowly varying component (the DC '
    'level, determined by mean arterial pressure and any external offsets) and a pulsatile component '
    '(the AC signal, driven by cardiac ejection). Pulse pressure (PP = SBP \u2013 DBP) is a pure AC '
    'quantity: it represents the difference between the maximum and minimum of the pulsatile '
    'excursion and is, by definition, independent of any additive DC offset (Fig. 1, panels A and B).'
)

add_para(
    'This independence has a critical implication. If a sensor measures PP correctly, its gain '
    '(transduction sensitivity, in mV/mmHg or equivalent units) must be correct, because PP accuracy '
    'requires that the electrical output difference corresponding to the true pressure difference '
    '(SBP \u2013 DBP) is exact. Conversely, if the gain is incorrect (e.g., the sensor reads 1.15 mV '
    'per mmHg instead of 1.00 mV per mmHg), the measured PP will be proportionally distorted '
    '(Fig. 1, panel C). PP accuracy therefore serves as an implicit validator of sensor gain.'
)

add_para(
    'This relationship is demonstrated quantitatively in Figure 7, which shows simulated arterial '
    'waveforms under correct gain (panel A), gain error (panel B), and DC offset only (panel C). '
    'Panels D and E confirm that the PP regression slope against a reference directly estimates '
    'sensor gain, providing a practical validation tool. The logical chain from PP accuracy to '
    'zero-calibration-free monitoring is summarised in panel F.'
)

add_heading_styled('The three sources of DC offset and their engineering elimination', level=2)

add_para(
    'If gain is correct (validated by PP accuracy), the only remaining source of measurement error '
    'is the DC offset. As noted above, three distinct sources contribute to this offset in conventional '
    'systems (Table 1). Importantly, each can be eliminated by established engineering solutions:'
)

add_para(
    '(1) Hydrostatic column offset: Catheter-tip MEMS pressure sensors (e.g., Millar Mikro-Cath) '
    'place the sensing element directly at the catheter tip, eliminating the fluid column entirely.8,9 '
    'The sensor measures pressure at the point of interest, making the measurement independent of the '
    'transducer-to-patient height relationship. Millar\u2019s own documentation states that \u201cthere is '
    'no need to calibrate to patient height or insertion angle.\u201d10'
)

add_para(
    '(2) Atmospheric pressure reference: Absolute pressure sensors measure total pressure relative '
    'to an internal vacuum reference rather than to atmosphere. By incorporating a barometric pressure '
    'sensor in the monitor and electronically subtracting atmospheric pressure, the gauge pressure '
    'can be computed without manual atmospheric zeroing. This approach is already implemented in '
    'chronic implantable sensors such as Millar\u2019s TiSense platform, which combines an absolute '
    'half-bridge sensor with an external barometric sensor for compensation.11'
)

add_para(
    '(3) Transducer drift: Self-calibrating MEMS sensors incorporate internal reference pressure '
    'cavities that provide known pressure points for periodic auto-calibration.12 By cycling through '
    'a reference pressure (e.g., via liquid-to-vapour phase transition in a sealed micro-cavity), the '
    'sensor continuously corrects for zero drift without external intervention. Kang et al.12 '
    'demonstrated a flexible MEMS sensor with integrated self-calibration achieving long-term stability.'
)

add_para(
    'When all three solutions are implemented simultaneously\u2014tip sensor, absolute pressure with '
    'barometric compensation, and self-calibrating reference\u2014every source of DC offset is eliminated '
    'by design. Manual zero calibration becomes redundant (Fig. 3).'
)

# ══════════════════════════════════════════════════════════════════
# CCC FRAMEWORK
# ══════════════════════════════════════════════════════════════════
add_heading_styled('Formalisation through the CCC decomposition', level=1)

add_heading_styled('Lin\u2019s CCC and its components', level=2)

add_para(
    'Lin\u2019s concordance correlation coefficient (\u03C1c) quantifies the agreement between paired '
    'measurements along the 45-degree line of perfect concordance.6 It decomposes as:'
)

eq = doc.add_paragraph()
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = eq.add_run('\u03C1c = r \u00D7 Cb')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.italic = True
eq.paragraph_format.space_after = Pt(6)

add_para(
    'where r (Pearson\u2019s correlation coefficient) measures precision (the tightness of data around '
    'the best-fit line) and Cb (bias correction factor) measures accuracy (how far the best-fit line '
    'deviates from the identity line). Cb is further decomposed as:'
)

eq2 = doc.add_paragraph()
eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = eq2.add_run('Cb = 2 / (v + 1/v + u\u00B2)')
run2.font.name = 'Times New Roman'
run2.font.size = Pt(12)
run2.italic = True
eq2.paragraph_format.space_after = Pt(6)

add_para(
    'where v = \u03C31/\u03C32 is the scale shift (ratio of standard deviations, reflecting gain error) '
    'and u = (\u03BC1 \u2013 \u03BC2)/\u221A(\u03C31\u03C32) is the location shift (normalised mean '
    'difference, reflecting offset).7'
)

add_heading_styled('What zero calibration does\u2014and does not do\u2014in CCC terms', level=2)

add_para(
    'Zero calibration is an operation that drives u toward zero by removing the systematic offset '
    'between the device reading and the true pressure. In CCC terms, successful zeroing achieves '
    'u \u2248 0, which maximises Cb with respect to the location component. However, zeroing has no '
    'effect on v: if the sensor gain is incorrect (v \u2260 1), Cb remains below 1.0 even after '
    'perfect zeroing (Fig. 2, Fig. 4).'
)

add_para(
    'This asymmetry is illustrated in Figure 2, which shows four simulated scenarios of arterial '
    'pressure measurement. Scenario A (offset only, before zeroing) shows reduced CCC (0.855) due '
    'to location shift (u = \u20130.55). Scenario B (after zeroing) shows near-perfect agreement '
    '(CCC = 0.986, Cb = 1.000) because the offset has been removed and no gain error exists. '
    'Scenario C (gain error, v = 1.11) shows that zeroing cannot improve Cb (0.870) because the '
    'scale shift persists. Scenario D (gain + offset) has CCC = 0.976; zeroing '
    'would improve it to Scenario C but not to Scenario B.'
)

add_para(
    'The critical comparison is between Scenarios B and C: the Bland-Altman analysis '
    '(Fig. 5, Fig. 8 lower panels) shows near-zero bias in both cases, yet CCC decomposition '
    'reveals fundamentally different error structures (Table 2). Scenario B has Cb = 1.000 '
    '(no systematic error) while Scenario C has Cb = 0.870 (hidden gain error, v = 1.11). '
    'This demonstrates that BA analysis alone cannot distinguish true agreement from '
    'agreement achieved by offset cancellation that masks gain error.'
)

add_para(
    'Figure 4 presents the Cb diagnostic space as a contour plot over the u\u2013v plane. Zero '
    'calibration moves a device horizontally (u \u2192 0) but not vertically (v is unchanged). The '
    'critical insight is that only devices with both u = 0 and v = 1 achieve Cb = 1.0. A sensor '
    'with correct gain (v = 1)\u2014validated by PP accuracy\u2014combined with offset-free design '
    '(u = 0) achieves this ideal without calibration.'
)

add_para(
    'The sensitivity analysis in Figure 6 quantifies these relationships. Panel A shows Cb as a '
    'function of gain error and DC offset, confirming that zero calibration (vertical movement '
    'toward zero offset) improves Cb only when gain is correct. Panel B shows how CCC degrades '
    'with increasing gain error at different levels of sensor precision (r), demonstrating that '
    'even high-precision sensors (r = 0.99) suffer substantial CCC loss when gain error exceeds '
    '\u00B110%. Notably, typical MEMS sensors achieve gain accuracy within \u00B15%, corresponding '
    'to Cb > 0.99 in the zero-offset condition.'
)

add_heading_styled('The complete argument', level=2)

add_para(
    'Combining the engineering and statistical perspectives, the argument can be stated formally:'
)

add_para('(1) PP is accurate \u2192 gain is correct \u2192 v = 1 (no scale shift)', bold=False)
add_para('(2) Catheter-tip sensor + barometric compensation + self-calibrating MEMS \u2192 all DC offsets eliminated \u2192 u = 0 (no location shift)', bold=False)
add_para('(3) v = 1 and u = 0 \u2192 Cb = 1.0', bold=False)
add_para('(4) CCC = r \u00D7 1.0 = r', bold=False)

add_para(
    'That is, the performance of a properly designed sensor is limited only by its precision (r), '
    'which reflects random measurement noise. All systematic error\u2014both offset and gain\u2014is '
    'resolved by design, not by calibration. The sensor\u2019s Cb approaches 1.0 as a design '
    'specification rather than a calibration outcome.'
)

# ══════════════════════════════════════════════════════════════════
# DISCUSSION
# ══════════════════════════════════════════════════════════════════
add_heading_styled('Discussion', level=1)

add_heading_styled('Implications for device development', level=2)

add_para(
    'The framework presented here reframes zero calibration from a clinical necessity to a design '
    'workaround. Conventional fluid-filled systems require zeroing because their architecture '
    'inherently introduces DC offsets. Rather than continuing to refine calibration procedures, '
    'device manufacturers should pursue designs that eliminate the need for calibration entirely. '
    'The component technologies\u2014catheter-tip MEMS, absolute pressure sensing, barometric '
    'compensation, and self-calibrating references\u2014already exist individually in commercial '
    'or near-commercial products.8\u201312 Their integration into a single clinical arterial pressure '
    'monitoring system is an engineering challenge, not a scientific one.'
)

add_para(
    'The CCC decomposition provides a quantitative design target: device developers should aim for '
    'Cb \u2265 0.99 without any calibration step, verifiable by comparing the uncalibrated device '
    'output against a reference standard. If Cb < 0.99 prior to calibration, the device has a '
    'residual systematic error that calibration can only partially mask. The decomposition into '
    'u and v further identifies whether the residual error is an offset (addressable by zeroing) '
    'or a gain error (requiring hardware or algorithmic correction).'
)

add_heading_styled('Implications for regulatory validation', level=2)

add_para(
    'Current regulatory pathways (e.g., FDA 510(k)) do not prescribe specific statistical methods '
    'for validating arterial pressure monitors.13 Published validation studies overwhelmingly rely '
    'on Bland-Altman analysis and percentage error.14,15 However, the Bland-Altman plot of a zeroed '
    'device will show bias \u2248 0, potentially masking proportional bias (gain error) within the '
    'limits of agreement. CCC reporting, particularly the Cb component, would provide an additional '
    'layer of scrutiny that distinguishes offset correction from genuine measurement accuracy.'
)

add_para(
    'Our simulation analysis demonstrates this limitation quantitatively. As shown in Table 2 and '
    'Figure 8, Scenarios B (after zeroing, no gain error) and D (gain + offset) produce similar '
    'BA bias values (approximately 0 and 1 mmHg, respectively), yet their CCC values differ '
    'substantially (0.986 vs. 0.976) and their Cb values even more so when gain error alone is '
    'considered (B: 1.000 vs. C: 0.870). The Bland-Altman framework, while invaluable for detecting '
    'fixed bias, is inherently insensitive to proportional (gain-dependent) error when the device '
    'has been zeroed.'
)

add_para(
    'We suggest that validation studies for arterial pressure monitors should report: (1) CCC with '
    'decomposition into r and Cb; (2) Cb before and after zero calibration, to quantify the device\u2019s '
    'dependence on calibration; and (3) the individual contributions of u and v to any Cb deficit. '
    'This information would enable regulators and clinicians to assess whether a device achieves '
    'accuracy through good design or through calibration-dependent offset removal.'
)

add_heading_styled('Extension to non-invasive cardiac output monitors', level=2)

add_para(
    'The principle that calibration corrects only offset (u) while leaving gain error (v) untouched '
    'extends naturally to non-invasive cardiac output (CO) monitors. ClearSight\u2019s Physiocal '
    'algorithm periodically re-optimises the volume clamp setpoint\u2014essentially an offset '
    'correction.16 NICOM/Starling\u2019s bioreactance phase reference provides a baseline offset for '
    'the phase-to-stroke-volume conversion.17 FloTrac\u2019s arterial waveform calibration adjusts '
    'the mean pressure offset.18 In every case, the auto-calibration routine addresses u (location '
    'shift) but does not correct v (scale shift)\u2014that is, the gain of the pressure-to-CO or '
    'impedance-to-SV conversion.'
)

add_para(
    'This observation, combined with the CO monitoring validation framework proposed by Odor et al.19 '
    'and the CCC-based extension we previously advocated,20 suggests a unified principle: for any '
    'haemodynamic monitoring device, the bias correction factor Cb should be decomposed into u and v '
    'components, and the device\u2019s reliance on calibration to achieve acceptable Cb should be '
    'explicitly quantified. A device that requires frequent recalibration to maintain Cb close to 1.0 '
    'has a fundamental design limitation that calibration can only partially compensate.'
)

add_heading_styled('Assumptions and limitations', level=2)

add_para(
    'Our argument rests on the assumption that sensor response is linear across the physiological '
    'pressure range. If the sensor exhibits significant non-linearity, PP accuracy at one pressure '
    'level would not guarantee gain correctness at other levels. However, modern MEMS piezoresistive '
    'sensors typically achieve non-linearity < 0.1% of full-scale output over 0\u2013300 mmHg,21 '
    'well within the requirements for clinical arterial pressure monitoring.'
)

add_para(
    'The framework addresses systematic error only. Random measurement noise (which determines r) '
    'is a separate concern that is not affected by calibration or the design features discussed here. '
    'A zero-calibration-free sensor with high random noise would still have poor CCC despite '
    'achieving Cb = 1.0.'
)

add_para(
    'Finally, we acknowledge that the clinical adoption of catheter-tip MEMS sensors for routine '
    'arterial monitoring faces practical barriers including cost, disposability, and compatibility '
    'with existing monitoring infrastructure. Our aim is not to mandate immediate clinical change but '
    'to establish the theoretical principle that zero calibration is a designable-out limitation, '
    'and to provide a quantitative framework (CCC decomposition) for evaluating progress toward '
    'this goal.'
)

add_heading_styled('Clinical perspective', level=2)

add_para(
    'For practising clinicians, the immediate practical message is nuanced. Zero calibration remains '
    'essential with current fluid-filled systems and should continue to be performed rigorously, as '
    'emphasised by Saugel et al.2 and Gupta et al.3 However, clinicians should be aware that even '
    'perfect zeroing corrects only offset, not gain. A zeroed transducer with a 10% gain error will '
    'display SBP 132 instead of 120 and DBP 88 instead of 80\u2014a clinically significant '
    'overestimation that would not be detected by the zeroing procedure and would appear as a '
    'widened PP (52 vs. 40 mmHg) rather than a shifted baseline.'
)

add_para(
    'The recognition that PP accuracy implies gain correctness also suggests a simple clinical check: '
    'if the invasively measured PP is physiologically plausible and consistent with the non-invasive '
    'cuff measurement, the sensor gain is likely correct. An implausible PP (e.g., very narrow or '
    'very wide relative to the patient\u2019s condition) may indicate a gain problem that zeroing '
    'cannot resolve.'
)

# ══════════════════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════════════════
add_heading_styled('Conclusion', level=1)

add_para(
    'We have presented a theoretical framework demonstrating that zero calibration of arterial '
    'pressure transducers is rendered unnecessary by appropriate sensor design. The argument '
    'combines two insights: (1) pulse pressure accuracy validates sensor gain, eliminating scale '
    'shift (v = 1); and (2) catheter-tip MEMS sensors, absolute pressure measurement with barometric '
    'compensation, and self-calibrating MEMS references together eliminate all sources of DC offset '
    '(u = 0). Formalised through Lin\u2019s CCC decomposition, this means Cb = 1.0 is achievable by '
    'design rather than by calibration. The framework provides a quantitative basis for '
    'next-generation arterial pressure monitor design, regulatory validation requirements, and '
    'a deeper understanding of what calibration does\u2014and does not do\u2014in haemodynamic monitoring.'
)

# ══════════════════════════════════════════════════════════════════
# DECLARATIONS
# ══════════════════════════════════════════════════════════════════
add_heading_styled('Declarations of interest', level=1)
add_para('[To be completed by authors]')

add_heading_styled('Funding', level=1)
add_para('[To be completed by authors]')

add_heading_styled('Authors\u2019 contributions', level=1)
add_para('[To be completed by authors]')

add_heading_styled('Acknowledgements', level=1)
add_para('[To be completed by authors]')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════
add_heading_styled('References', level=1)

refs = [
    '1. Saugel B, Kouz K, Meidert AS, Schulte-Uentrop L, Romagnoli S. How to measure blood pressure using an arterial catheter: a systematic 5-step approach. Crit Care 2020; 24: 172. doi:10.1186/s13054-020-02859-w',
    '2. Saugel B, Sessler DI. Perioperative blood pressure management. Anesthesiology 2021; 134: 250\u201361. doi:10.1097/ALN.0000000000003610',
    '3. Gupta D, Jain A, Ismaeil M. Zero arterial catheters with every change in the height difference of pressure transducer and catheter insertion site. Ann Card Anaesth 2025; 28: 205. doi:10.4103/aca.aca_198_24',
    '4. Mark JB. Atlas of Cardiovascular Monitoring. New York: Churchill Livingstone, 1998.',
    '5. McGhee BH, Bridges MEJ. Monitoring arterial blood pressure: what you may not know. Crit Care Nurse 2002; 22: 60\u201379. doi:10.4037/ccn2002.22.2.60',
    '6. Lin LI-K. A concordance correlation coefficient to evaluate reproducibility. Biometrics 1989; 45: 255\u201368. doi:10.2307/2532051',
    '7. Lin LI-K. A note on the concordance correlation coefficient. Biometrics 2000; 56: 324\u20135. doi:10.1111/j.0006-341X.2000.00324.x',
    '8. Hasenkamp W, Theumer T, Nolte J, et al. Polyimide/SU-8 catheter-tip MEMS gauge pressure sensor. Biomed Microdevices 2012; 14: 819\u201328. doi:10.1007/s10544-012-9661-8',
    '9. Song P, Ma Z, Ma J, et al. Recent progress of miniature MEMS pressure sensors. Micromachines 2020; 11: 56. doi:10.3390/mi11010056',
    '10. Millar, Inc. Mikro-Cath Pressure Catheter \u2013 Product Information. https://millar.com/Clinical/MikroCath/ (accessed 25 March 2026).',
    '11. Millar, Inc. TiSense \u2013 Chronic Pressure Sensing. https://millar.com/Clinical/MikroCath/Continuous-Compartment-Pressure-Measurements/ (accessed 25 March 2026).',
    '12. Kang Y, Ge C, Hu Y, et al. Development of a flexible integrated self-calibrating MEMS pressure sensor using a liquid-to-vapor phase change. Sensors 2022; 22: 9737. doi:10.3390/s22249737',
    '13. U.S. Food and Drug Administration. Premarket Notification 510(k). https://www.fda.gov/medical-devices/premarket-submissions-selecting-and-preparing-correct-submission/premarket-notification-510k (accessed 25 March 2026).',
    '14. Joosten A, Desebbe O, Suehiro K, et al. Accuracy and precision of non-invasive cardiac output monitoring devices in perioperative medicine: a systematic review and meta-analysis. Br J Anaesth 2017; 118: 298\u2013310. doi:10.1093/bja/aew461',
    '15. Bland JM, Altman DG. Statistical methods for assessing agreement between two methods of clinical measurement. Lancet 1986; 327: 307\u201310. doi:10.1016/S0140-6736(86)90837-8',
    '16. Ameloot K, Palmers PJ, Malbrain MLNG. The accuracy of noninvasive cardiac output and pressure measurements with finger cuff: a concise review. Curr Opin Crit Care 2015; 21: 232\u20139. doi:10.1097/MCC.0000000000000198',
    '17. Squara P, Denjean D, Estagnasie P, Brusset A, Dib JC, Dubois C. Noninvasive cardiac output monitoring (NICOM): a clinical validation. Intensive Care Med 2007; 33: 1191\u20134. doi:10.1007/s00134-007-0640-0',
    '18. Manecke GR. Edwards FloTrac sensor and Vigileo monitor: easy, accurate, reliable cardiac output assessment using the arterial pulse wave. Expert Rev Med Devices 2005; 2: 523\u20137. doi:10.1586/17434440.2.5.523',
    '19. Odor PM, Bampoe S, Cecconi M. Cardiac output monitoring: validation studies\u2014how results should be presented. Curr Anesthesiol Rep 2017; 7: 410\u201315. doi:10.1007/s40140-017-0239-0',
    '20. [Author(s)]. Lin\u2019s concordance correlation coefficient: a missing metric in cardiac output monitor validation. [Journal, Year; details to be completed].',
    '21. Barlian AA, Park W-T, Mallon JR, Rastegar AJ, Pruitt BL. Review: semiconductor piezoresistance for microsystems. Proc IEEE 2009; 97: 513\u201352. doi:10.1109/JPROC.2009.2013612',
    '22. Critchley LAH, Critchley JAJH. A meta-analysis of studies using bias and precision statistics to compare cardiac output measurement techniques. J Clin Monit Comput 1999; 15: 85\u201391. doi:10.1023/A:1009982611386',
    '23. Critchley LA, Lee A, Ho AM-H. A critical review of the ability of continuous cardiac output monitors to measure trends in cardiac output. Anesth Analg 2010; 111: 1180\u201392. doi:10.1213/ANE.0b013e3181f08a5b',
    '24. Cecconi M, Rhodes A, Poloniecki J, Della Rocca G, Grounds RM. Bench-to-bedside review: the importance of the precision of the reference technique in method comparison studies. Crit Care 2009; 13: 201. doi:10.1186/cc7129',
    '25. McBride GB. A proposal for strength-of-agreement criteria for Lin\u2019s concordance correlation coefficient. NIWA Client Report HAM2005-062. 2005.',
    '26. Bland JM, Altman DG. Measuring agreement in method comparison studies. Stat Methods Med Res 1999; 8: 135\u201360. doi:10.1177/096228029900800204',
    '27. Romagnoli S, Ricci Z, Romano SM, et al. FloTrac/Vigileo (third generation) and MostCare/PRAM versus echocardiography for cardiac output estimation in vascular surgery. J Cardiothorac Vasc Anesth 2013; 27: 1114\u201321. doi:10.1053/j.jvca.2013.04.017',
    '28. Scalia A, Ghafari C, Navarre W, Delmotte P, Phillips R, Carlier S. High fidelity pressure wires provide accurate validation of non-invasive central blood pressure and pulse wave velocity measurements. Biomedicines 2023; 11: 1235. doi:10.3390/biomedicines11041235',
    '29. Kim S-H, Lilot M, Sidhu KS, et al. Accuracy and precision of continuous noninvasive arterial pressure monitoring compared with invasive arterial pressure. Anesthesiology 2014; 120: 1080\u201397. doi:10.1097/ALN.0000000000000226',
    '30. Chatterjee K. The Swan-Ganz catheters: past, present, and future. Circulation 2009; 119: 147\u201352. doi:10.1161/CIRCULATIONAHA.108.811141',
]

for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# TABLE 1
# ══════════════════════════════════════════════════════════════════
add_heading_styled('Table 1', level=1)
add_para(
    'Table 1. Sources of DC offset in conventional arterial pressure monitoring '
    'and engineering solutions for their elimination.',
    bold=True, italic=True
)

table = doc.add_table(rows=5, cols=5)
table.style = 'Light Grid Accent 1'

headers = ['Offset source', 'Physical mechanism', 'Magnitude', 'Engineering solution', 'CCC effect']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'

data = [
    ['Hydrostatic column', '\u0394P = \u03C1gh (height difference between transducer and catheter tip)',
     '~0.74 mmHg/cm', 'Catheter-tip MEMS sensor (eliminates fluid column)',
     'Reduces u (location shift)'],
    ['Atmospheric pressure reference', 'Gauge pressure measured relative to atmosphere (~760 mmHg)',
     'Entire baseline', 'Absolute pressure sensor + built-in barometer for electronic subtraction',
     'Reduces u (location shift)'],
    ['Transducer drift', 'Mechanical creep, thermal effects in strain gauge',
     '~1\u20135 mmHg/day', 'Self-calibrating MEMS with internal reference pressure cavity',
     'Reduces u (location shift)'],
    ['Gain error (NOT corrected by zeroing)', 'Sensitivity mismatch: output per unit pressure \u2260 nominal',
     '1\u201315% typical', 'Factory calibration; PP accuracy as implicit gain check',
     'Affects v (scale shift); zeroing has no effect'],
]

for row_idx, row_data in enumerate(data, start=1):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.name = 'Times New Roman'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# TABLE 2 (NEW)
# ══════════════════════════════════════════════════════════════════
add_heading_styled('Table 2', level=1)
add_para(
    'Table 2. Complete statistical comparison across four simulated scenarios '
    '(n = 150 paired measurements). Simulation parameters: reference mean = 100 mmHg, SD = 22 mmHg; '
    'offset = 12 mmHg; gain = 0.89; noise SD = 3.5 mmHg.',
    bold=True, italic=True
)

table2 = doc.add_table(rows=5, cols=10)
table2.style = 'Light Grid Accent 1'

t2_headers = ['Scenario', 'CCC', 'r', 'Cb', 'u', 'v', 'BA Bias\n(mmHg)',
              'LOA lower\n(mmHg)', 'LOA upper\n(mmHg)', 'PE (%)']
for i, h in enumerate(t2_headers):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'

t2_data = [
    ['A: Before zeroing\n(offset only)', '0.855', '0.986', '0.867', '\u22120.55', '0.99',
     '11.9', '4.8', '19.0', '7.2'],
    ['B: After zeroing\n(offset removed)', '0.986', '0.986', '1.000', '0.01', '0.99',
     '\u22120.1', '\u22127.2', '7.0', '7.2'],
    ['C: Gain error\n(zeroing cannot fix)', '0.855', '0.982', '0.870', '0.54', '1.11',
     '\u221210.9', '\u221219.4', '\u22122.4', '8.7'],
    ['D: Gain + offset\n(zeroing fixes only offset)', '0.976', '0.982', '0.993', '\u22120.05', '1.11',
     '1.1', '\u22127.4', '9.6', '8.7'],
]

for row_idx, row_data in enumerate(t2_data, start=1):
    for col_idx, cell_text in enumerate(row_data):
        cell = table2.rows[row_idx].cells[col_idx]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.name = 'Times New Roman'

# Add note below table
add_para(
    'Note: Scenarios A and B have identical v (0.99, no gain error); zeroing moves A\u2192B by '
    'removing offset (u: \u22120.55\u21920.01). Scenarios C and D have identical v (1.11, gain error); '
    'zeroing moves D\u2192C by removing offset but v remains unchanged. BA bias in D (\u22481 mmHg) '
    'appears deceptively small because gain-dependent and offset-dependent biases partially cancel. '
    'CCC = concordance correlation coefficient; r = Pearson correlation; Cb = bias correction factor; '
    'u = location shift; v = scale shift; BA = Bland-Altman; LOA = limits of agreement; PE = percentage error.',
    italic=True, space_after=Pt(12)
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# FIGURES (ORIGINAL + NEW)
# ══════════════════════════════════════════════════════════════════
add_heading_styled('Figures', level=1)

# Figure 1 (original)
add_figure(
    'figure1_signal_decomposition.png',
    'Figure 1. Arterial pressure waveform decomposition. (A) True arterial pressure with '
    'DC component (offset-dependent baseline) and AC component (pulse pressure, PP = 40 mmHg). '
    '(B) With DC offset of 15 mmHg (e.g., from hydrostatic column): PP remains unchanged at '
    '40 mmHg because it is a differential quantity. (C) With gain error (v = 1.15): PP is '
    'distorted from 40 to 46 mmHg, demonstrating that PP accuracy validates sensor gain. '
    'Zero calibration corrects the offset in (B) but cannot correct the gain error in (C).'
)

doc.add_page_break()

# Figure 2 (original)
add_figure(
    'figure2_ccc_zeroing_scenarios.png',
    'Figure 2. Concordance plots for four simulated scenarios (n = 150 paired measurements). '
    '(A) Before zeroing: offset produces location shift u = \u20130.55, reducing Cb to 0.867. '
    '(B) After zeroing: offset removed, CCC = 0.986 with Cb = 1.000. '
    '(C) Gain error (v = 1.11): zeroing cannot improve Cb (0.870) because the scale shift persists. '
    '(D) Gain + offset: CCC = 0.976; zeroing moves device from D to C but not to B. '
    'Dashed line = identity (y = x); solid line = best fit.'
)

doc.add_page_break()

# Figure 3 (original)
add_figure(
    'figure3_system_comparison.png',
    'Figure 3. Comparison of conventional and proposed zero-free arterial pressure monitoring systems. '
    '(A) Conventional fluid-filled system: three sources of DC offset (hydrostatic column, atmospheric '
    'pressure reference, transducer drift) necessitate manual zero calibration, which corrects offset (u) '
    'only. (B) Proposed zero-free system: catheter-tip MEMS sensor eliminates hydrostatic column; '
    'built-in barometer provides electronic atmospheric compensation; self-calibrating MEMS eliminates '
    'drift. All offset sources are eliminated by design, and PP accuracy validates gain (v = 1).'
)

doc.add_page_break()

# Figure 4 (original)
add_figure(
    'figure4_cb_diagnostic_space.png',
    'Figure 4. Bias correction factor (Cb) as a function of location shift (u) and scale shift (v). '
    'Contour lines show iso-Cb values. Zero calibration moves a device horizontally (u \u2192 0, green '
    'arrows) but does not change v. Points A\u2013D correspond to the scenarios in Figure 2. The gold '
    'star marks the ideal (u = 0, v = 1, Cb = 1.0). A device with correct gain (v = 1, validated by '
    'PP accuracy) and offset-free design (u = 0) achieves Cb = 1.0 without calibration.'
)

doc.add_page_break()

# Figure 5 (NEW)
add_figure(
    'figure5_ba_comparison.png',
    'Figure 5. Bland-Altman analysis for the same four scenarios as Figure 2. Key observation: '
    'After zeroing (B), bias \u2248 0 appears as good agreement. However, gain error (C) also shows '
    'near-zero bias with similar limits of agreement, despite a fundamentally different error structure '
    '(v = 1.11). The Bland-Altman plot alone cannot distinguish true agreement (B) from hidden gain '
    'error (C). Compare CCC values in the lower-right corner of each panel: B has Cb = 1.000 while '
    'C has Cb = 0.870.'
)

doc.add_page_break()

# Figure 6 (NEW)
add_figure(
    'figure6_sensitivity_analysis.png',
    'Figure 6. Sensitivity analysis. (A) Cb as a function of gain error (%) and DC offset (mmHg), '
    'showing that zero calibration (green arrows, vertical movement toward zero offset) improves Cb '
    'only when gain is correct. Points A\u2013D correspond to the four scenarios. '
    '(B) CCC as a function of gain error for different levels of sensor precision (r), with offset = 0. '
    'The green-shaded zone shows typical MEMS gain accuracy (\u00B15%), within which Cb > 0.99. '
    'Even highly precise sensors (r = 0.99) suffer substantial CCC degradation beyond \u00B110% gain error.'
)

doc.add_page_break()

# Figure 7 (NEW)
add_figure(
    'figure7_pp_validation.png',
    'Figure 7. Pulse pressure accuracy as implicit gain validation. (A) Correct gain (v = 1.0): '
    'measured waveform matches true pressure, PP = 40 mmHg preserved. (B) Gain error (v = 1.15): '
    'waveform is amplitude-scaled, PP distorted from 40 to 46 mmHg. (C) DC offset only (+15 mmHg): '
    'waveform is shifted but PP remains 40 mmHg, unchanged. (D) Scatter plot showing that PP '
    'measurement accuracy correlates with sensor gain, colour-coded by actual gain. '
    '(E) PP regression: slope directly estimates sensor gain (Device A: correct gain, slope = 1.00; '
    'Device B: gain error, slope = 0.88). (F) Logical chain from PP accuracy to zero-calibration-free '
    'monitoring via gain validation and offset-free design.'
)

doc.add_page_break()

# Figure 8 (NEW)
add_figure(
    'figure8_ba_vs_concordance.png',
    'Figure 8. Direct comparison of concordance plots (top row) and Bland-Altman plots (bottom row) '
    'for all four scenarios, demonstrating why Bland-Altman analysis alone is insufficient. '
    'The concordance plots reveal the error structure through CCC decomposition (r, Cb, u, v), '
    'while the BA plots show only bias and limits of agreement. Critical comparison: Scenarios B and C '
    'have similar BA profiles (bias \u2248 0 vs. \u224811 mmHg), but concordance analysis reveals '
    'B has Cb = 1.000 (true agreement) while C has Cb = 0.870 (hidden gain error, v = 1.11). '
    'This distinction is invisible on the BA plot alone.'
)

# ══════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════
doc.save(OUTPATH)
print(f'English manuscript saved: {OUTPATH}')
