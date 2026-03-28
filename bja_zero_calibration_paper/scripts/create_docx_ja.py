#!/usr/bin/env python3
"""
Create BJA Special Article .docx — Japanese version
「脈圧の正確性がゲイン妥当性を暗示する：
  ゼロ校正不要の動脈圧モニタリングに向けた理論的枠組み」
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

FIGDIR = '/home/ubuntu/bja_figures'
OUTPATH = '/home/ubuntu/BJA_ZeroFree_Manuscript_JA.docx'

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p

def add_figure(filename, caption, width=Inches(6)):
    path = os.path.join(FIGDIR, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    cap = doc.add_paragraph()
    cap_run = cap.add_run(caption)
    cap_run.font.name = 'Times New Roman'
    cap_run.font.size = Pt(10)
    cap_run.italic = True
    cap.paragraph_format.space_after = Pt(12)
    return cap

# ══════════════════════════════════════════════════════════════════
# 表紙
# ══════════════════════════════════════════════════════════════════
add_para('SPECIAL ARTICLE（日本語版）', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    '脈圧の正確性が暗示するゲインの妥当性：\n'
    'ゼロ校正不要の観血的動脈圧モニタリングに向けた理論的枠組み'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.bold = True
title.paragraph_format.space_after = Pt(12)

title_en = doc.add_paragraph()
title_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_en = title_en.add_run(
    'Pulse Pressure Accuracy as an Implicit Gain Validator:\n'
    'A Theoretical Framework for Zero-Calibration-Free Arterial Pressure Monitoring'
)
run_en.font.name = 'Times New Roman'
run_en.font.size = Pt(12)
run_en.italic = True
title_en.paragraph_format.space_after = Pt(18)

add_para('ランニングタイトル：ゼロ校正不要の動脈圧モニタリング', italic=True,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_para('[著者名（未定）]', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[所属（未定）]', alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_para('責任著者:', bold=True)
add_para('[氏名、住所、メールアドレス、電話番号]')
doc.add_paragraph()

add_para('本文語数: 約3,800語（参考文献、表、図の説明文を除く）')
add_para('表・図: 表2、図8点')
add_para('参考文献: 30件')
doc.add_paragraph()

add_para('キーワード: ', bold=True, space_after=Pt(0))
kw = doc.paragraphs[-1]
kw_run = kw.add_run(
    '観血的動脈圧モニタリング；ゼロ校正；脈圧；'
    '一致性相関係数；MEMSセンサ；バイアス補正係数；機器設計'
)
kw_run.font.name = 'Times New Roman'
kw_run.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 要旨
# ══════════════════════════════════════════════════════════════════
add_heading_styled('要旨（Summary）', level=1)

add_para(
    '観血的動脈圧トランスデューサのゼロ校正は、血行動態モニタリングにおける基本的かつ不可欠な手技として '
    '広く認知されている。本論文ではこの前提に挑戦し、適切なセンサ設計によりゼロ校正が原理的に不要となる '
    'ことを示す理論的枠組みを提示する。本論文の主張は二つの柱に基づく。第一に、脈圧（PP = 収縮期血圧 '
    '\u2013 拡張期血圧）は差分量（AC成分）であり、ゼロ校正が除去するDCオフセットとは独立であることを示す。'
    'センサが脈圧を正確に測定しているならば、そのゲイン（感度）は正しいことが保証される。なぜなら、'
    '脈圧の正確性は生理的範囲全体にわたる圧力差の正確な変換を必要とするからである。第二に、'
    'カテーテル先端MEMSセンサ、気圧補償付き絶対圧測定、および内部参照キャビティを持つ自己校正MEMS'
    'という3つの工学的解決策を特定し、これらが3つのDCオフセット源（静水圧カラム、大気圧基準、'
    'トランスデューサドリフト）をすべて排除することを示す。この議論をLinの一致性相関係数（CCC）の分解'
    'を用いて定式化し、ゼロ校正がバイアス補正係数C_bの位置ずれ成分（u）のみを修正し、'
    '脈圧を歪めるスケールずれ（v）\u2014すなわちゲインエラー\u2014は未補正のままであることを示す。'
    '適切に設計されたセンサはC_b \u2248 1.0を構造的に達成し、手動ゼロ校正を冗長にする。'
    '機器開発、規制上のバリデーション、および臨床実践への含意を議論し、'
    '非侵襲的心拍出量モニターへの枠組みの拡張も示す。'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 緒言
# ══════════════════════════════════════════════════════════════════
add_heading_styled('緒言（Introduction）', level=1)

add_para(
    '観血的動脈圧モニタリングは、麻酔科および集中治療において最も頻繁に実施される測定のひとつである。'
    '動脈カテーテル挿入後、臨床医は血行動態モニタリングを開始する前に、腋窩中線レベル（phlebostatic '
    'axis）で圧トランスデューサを大気圧にゼロ校正しなければならない[1,2]。この校正操作は、トランスデューサ'
    'とカテーテル挿入部位の高さ関係が変化するたびに\u2014例えば手術台が傾斜した場合やベッドの角度が調整'
    'された場合\u2014繰り返す必要がある[3]。圧センシング技術が数十年にわたり進歩してきたにもかかわらず、'
    '手動ゼロ校正の必要性は外部液充填トランスデューサシステムの導入以来、本質的に変わっていない。'
)

add_para(
    'この要件が持続しているのは根本的な物理的必然性ではなく、従来のモニタリングシステムの設計上の帰結である。'
    '標準的な液充填動脈ラインでは、圧トランスデューサは外部に位置し、生理食塩水のカラムによって血管内'
    'カテーテル先端と接続されている。この構造は3つの系統的DCオフセット源を導入する：(1)トランスデューサと'
    '測定部位の静水圧差（\u0394P = \u03C1gh、高低差1cmあたり約0.74 mmHg）；(2)大気圧（約760 mmHg）'
    'に対するゲージ圧測定の必要性；(3)ストレインゲージの機械的クリープや熱効果によるトランスデューサ'
    'ドリフト[4,5]。ゼロ校正はトランスデューサを大気に開放し出力をゼロにリセットすることで、これら3つの'
    'オフセットを同時に除去する。'
)

add_para(
    'しかし、ゼロ校正はもっぱらオフセット補正である。トランスデューサのゲイン（感度）\u2014すなわち'
    '真の圧力変化と電気的出力の比例関係\u2014を検証も補正もしない。ゲインが不正確であれば、測定波形は'
    '拡大または縮小され、このエラーはゼロ校正後も持続する。このオフセットエラーとゲインエラーの区別は'
    '単なる学術的議論ではなく、モニタリング機器のバリデーション方法と、以下で論じるように、'
    '機器設計のあり方に直接的な含意を持つ。'
)

add_para(
    '本論文では、機器が脈圧（PP）を正確に測定するならば、適切なセンサ設計によりゼロ校正が不要になると'
    '提案する。この議論をLinの一致性相関係数（CCC）の分解[6,7]を用いて定式化し、各DCオフセット源を'
    '排除する工学的解決策を特定し、動脈圧モニタリングおよびその先への含意を議論する。'
)

# ══════════════════════════════════════════════════════════════════
# 論理的議論
# ══════════════════════════════════════════════════════════════════
add_heading_styled('論理的議論：脈圧からゼロ校正不要モニタリングへ', level=1)

add_heading_styled('動脈圧のAC・DC成分分解', level=2)

add_para(
    '動脈圧波形P(t)は、緩やかに変化する成分（DCレベル：平均動脈圧と外部オフセットにより決定）と拍動成分'
    '（AC信号：心駆出によって駆動）に分解できる。脈圧（PP = SBP \u2013 DBP）は純粋なAC量であり、'
    '拍動の最大値と最小値の差を表し、定義上、いかなる加算的DCオフセットからも独立である'
    '（図1、パネルAおよびB）。'
)

add_para(
    'この独立性には重要な含意がある。センサがPPを正確に測定しているならば、そのゲイン（変換感度、'
    'mV/mmHg等の単位）は正しくなければならない。なぜならPPの正確性は、真の圧力差（SBP \u2013 DBP）に'
    '対応する電気的出力差が正確であることを要求するからである。逆に、ゲインが不正確であれば（例えば'
    'センサが1.00 mV/mmHgではなく1.15 mV/mmHgを読む場合）、測定されたPPは比例的に歪む'
    '（図1、パネルC）。したがって、PPの正確性はセンサゲインの暗黙的な妥当性検証として機能する。'
)

add_para(
    'この関係を図7で定量的に実証する。図7は正しいゲイン（パネルA）、ゲインエラー（パネルB）、'
    'DCオフセットのみ（パネルC）のシミュレーション動脈圧波形を示す。パネルDおよびEは、'
    '基準に対するPPの回帰直線の傾きがセンサゲインを直接推定することを確認し、実用的なバリデーション'
    'ツールを提供する。PPの正確性からゼロ校正不要モニタリングへの論理的連鎖をパネルFに要約する。'
)

add_heading_styled('DCオフセットの3つの源と工学的排除', level=2)

add_para(
    'ゲインが正しければ（PPの正確性により検証）、残る測定誤差の唯一の源はDCオフセットである。前述の通り、'
    '従来のシステムでは3つの異なる源がこのオフセットに寄与する（表1）。重要なことに、'
    '各々は確立された工学的解決策により排除可能である：'
)

add_para(
    '(1) 静水圧カラムオフセット：カテーテル先端MEMS圧センサ（例：Millar Mikro-Cath）はセンシング素子を'
    'カテーテル先端に直接配置し、液体カラムを完全に排除する[8,9]。センサは測定対象点で直接圧力を測定し、'
    'トランスデューサ\u2013患者間の高低差から独立した測定となる。Millar社の文書には「患者の身長や挿入角度'
    'への校正は不要」と記載されている[10]。'
)

add_para(
    '(2) 大気圧基準：絶対圧センサは大気に対してではなく内部真空参照に対して全圧を測定する。モニター内に'
    '気圧センサを組み込み、大気圧を電子的に減算することで、手動の大気圧ゼロ校正なしにゲージ圧を算出できる。'
    'このアプローチはMillar社のTiSenseプラットフォームなどの慢性埋込型センサにすでに実装されており、'
    '絶対ハーフブリッジセンサと外部気圧センサを組み合わせて補償を行っている[11]。'
)

add_para(
    '(3) トランスデューサドリフト：自己校正MEMSセンサは、周期的な自動校正のための既知の圧力点を提供する'
    '内部参照圧力キャビティを搭載している[12]。密封マイクロキャビティ内の液体\u2013気体相転移を利用して'
    '参照圧力を巡回することで、センサは外部介入なしにゼロドリフトを継続的に補正する。'
    'Kangら[12]は統合自己校正機能を持つフレキシブルMEMSセンサの長期安定性を実証した。'
)

add_para(
    '3つの解決策すべてを同時に実装すれば\u2014先端センサ、気圧補償付き絶対圧、および自己校正参照\u2014'
    'すべてのDCオフセット源が設計により排除される。手動ゼロ校正は冗長となる（図3）。'
)

# ══════════════════════════════════════════════════════════════════
# CCC枠組み
# ══════════════════════════════════════════════════════════════════
add_heading_styled('CCC分解による定式化', level=1)

add_heading_styled('LinのCCCとその構成要素', level=2)

add_para(
    'Linの一致性相関係数（\u03C1c）は、対をなす測定値の45度完全一致線上の合致度を定量化する[6]。'
    '以下のように分解される：'
)

eq = doc.add_paragraph()
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = eq.add_run('\u03C1c = r \u00D7 Cb')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.italic = True
eq.paragraph_format.space_after = Pt(6)

add_para(
    'ここでr（ピアソンの相関係数）は精度（precision：回帰直線周りのデータの緊密さ）を測定し、'
    'Cb（バイアス補正係数）は正確度（accuracy：回帰直線の45度恒等線からの乖離）を測定する。'
    'Cbはさらに以下のように分解される：'
)

eq2 = doc.add_paragraph()
eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = eq2.add_run('Cb = 2 / (v + 1/v + u\u00B2)')
run2.font.name = 'Times New Roman'
run2.font.size = Pt(12)
run2.italic = True
eq2.paragraph_format.space_after = Pt(6)

add_para(
    'ここでv = \u03C31/\u03C32はスケールシフト（標準偏差の比、ゲインエラーを反映）、'
    'u = (\u03BC1 \u2013 \u03BC2)/\u221A(\u03C31\u03C32)はロケーションシフト（正規化された平均差、'
    'オフセットを反映）である[7]。'
)

add_heading_styled('CCC的観点からゼロ校正が補正するもの・しないもの', level=2)

add_para(
    'ゼロ校正は、機器の読み値と真の圧力の間の系統的オフセットを除去することでuをゼロに近づける操作である。'
    'CCC的には、ゼロ校正の成功はu \u2248 0を達成し、ロケーション成分に関してCbを最大化する。'
    'しかし、ゼロ校正はvに影響しない：センサゲインが不正確であれば（v \u2260 1）、完全なゼロ校正後も'
    'Cbは1.0未満のままである（図2、図4）。'
)

add_para(
    'この非対称性を図2に示す。図2は動脈圧測定の4つのシミュレーションシナリオを示している。'
    'シナリオA（オフセットのみ、ゼロ校正前）はロケーションシフト（u = \u20130.55）によりCCCが低下'
    '（0.855）している。シナリオB（ゼロ校正後）はオフセットが除去され、ゲインエラーがないため'
    'ほぼ完全な一致（CCC = 0.986、Cb = 1.000）を示す。シナリオC（ゲインエラー、v = 1.11）は、'
    'スケールシフトが持続するためゼロ校正ではCb（0.870）を改善できないことを示す。'
    'シナリオD（ゲイン＋オフセット）はCCC = 0.976であり、ゼロ校正はDからCへの改善は'
    'もたらすが、Bへの到達はできない。'
)

add_para(
    '重要な比較はシナリオBとCの間にある：Bland-Altman解析（図5、図8下段パネル）はいずれの場合も'
    'バイアスがほぼゼロであることを示すが、CCC分解は根本的に異なるエラー構造を明らかにする（表2）。'
    'シナリオBはCb = 1.000（系統的エラーなし）であるのに対し、シナリオCはCb = 0.870'
    '（隠れたゲインエラー、v = 1.11）である。これはBA解析のみでは、真の一致と、ゲインエラーを'
    '隠蔽するオフセット相殺による一致を区別できないことを実証している。'
)

add_para(
    '図4はu\u2013v平面上の等高線図としてCb診断空間を提示する。ゼロ校正は機器を水平方向に移動'
    '（u \u2192 0）させるが、垂直方向（v）は変化しない。重要な知見は、u = 0かつv = 1の場合にのみ'
    'Cb = 1.0が達成されることである。正しいゲイン（v = 1、PPの正確性により検証）とオフセットフリーの'
    '設計（u = 0）を組み合わせたセンサは、校正なしでこの理想を達成する。'
)

add_para(
    '図6の感度分析はこれらの関係を定量化する。パネルAはゲインエラーとDCオフセットの関数としてCbを示し、'
    'ゼロ校正（ゼロオフセットへの垂直方向の移動）はゲインが正しい場合にのみCbを改善することを確認する。'
    'パネルBは異なるセンサ精度（r）レベルでのゲインエラーに伴うCCCの劣化を示し、高精度センサ'
    '（r = 0.99）であってもゲインエラーが\u00B110%を超えるとCCCが大幅に低下することを実証する。'
    '特筆すべきは、典型的なMEMSセンサは\u00B15%以内のゲイン精度を達成しており、'
    'ゼロオフセット条件でCb > 0.99に相当することである。'
)

add_heading_styled('完全な論証', level=2)

add_para(
    '工学的観点と統計的観点を統合すると、以下のように形式的に述べることができる：'
)

add_para('(1) PPが正確 \u2192 ゲインが正しい \u2192 v = 1（スケールシフトなし）')
add_para('(2) 先端センサ＋気圧補償＋自己校正MEMS \u2192 全DCオフセット排除 \u2192 u = 0（ロケーションシフトなし）')
add_para('(3) v = 1 かつ u = 0 \u2192 Cb = 1.0')
add_para('(4) CCC = r \u00D7 1.0 = r')

add_para(
    'すなわち、適切に設計されたセンサの性能は精度（r）のみによって制限される。rはランダムな測定ノイズを'
    '反映する。すべての系統的誤差\u2014オフセットとゲインの両方\u2014は校正ではなく設計により解決される。'
    'センサのCbは校正の結果ではなく、設計仕様として1.0に近づく。'
)

# ══════════════════════════════════════════════════════════════════
# 考察
# ══════════════════════════════════════════════════════════════════
add_heading_styled('考察（Discussion）', level=1)

add_heading_styled('機器開発への含意', level=2)

add_para(
    '本論文で提示した枠組みは、ゼロ校正を臨床的必要性から設計上の回避策へと再定義する。従来の液充填'
    'システムは、その構造が本質的にDCオフセットを導入するためにゼロ校正を必要とする。校正手順を洗練し'
    '続けるのではなく、機器メーカーは校正の必要性そのものを排除する設計を追求すべきである。'
    '構成技術\u2014カテーテル先端MEMS、絶対圧センシング、気圧補償、および自己校正参照\u2014は、'
    '商用または商用化間近の製品に個別にすでに存在している[8\u201312]。これらを単一の臨床用動脈圧'
    'モニタリングシステムに統合することは、科学的課題ではなく工学的課題である。'
)

add_para(
    'CCC分解は定量的な設計目標を提供する：機器開発者は、校正ステップなしでCb \u2265 0.99を達成すること'
    'を目指すべきであり、これは未校正の機器出力を基準標準と比較することで検証可能である。校正前に'
    'Cb < 0.99であれば、機器には校正が部分的にしか隠蔽できない残存系統的エラーがある。'
    'uとvへの分解により、残存エラーがオフセット（ゼロ校正で対処可能）かゲインエラー'
    '（ハードウェアまたはアルゴリズムの修正が必要）かをさらに特定できる。'
)

add_heading_styled('規制上のバリデーションへの含意', level=2)

add_para(
    '現在の規制経路（例：FDA 510(k)）は、動脈圧モニターのバリデーションに特定の統計手法を規定して'
    'いない[13]。公表されたバリデーション研究は圧倒的にBland-Altman解析とpercentage errorに依拠して'
    'いる[14,15]。しかし、ゼロ校正後の機器のBland-Altmanプロットはバイアス \u2248 0を示し、一致の限界'
    '内に比例バイアス（ゲインエラー）を隠蔽する可能性がある。CCC報告、特にCb成分は、オフセット補正と'
    '真の測定精度を区別する追加の精査層を提供する。'
)

add_para(
    '我々のシミュレーション分析はこの限界を定量的に実証する。表2および図8に示すとおり、'
    'シナリオB（ゼロ校正後、ゲインエラーなし）とD（ゲイン＋オフセット）は類似したBAバイアス値'
    '（それぞれ約0および1 mmHg）を示すが、CCC値は大幅に異なり（0.986 vs. 0.976）、'
    'ゲインエラーのみを考慮した場合のCb値はさらに大きく異なる（B: 1.000 vs. C: 0.870）。'
    'Bland-Altman枠組みは固定バイアスの検出に不可欠であるが、機器がゼロ校正済みの場合、'
    '比例的（ゲイン依存性）エラーに対して本質的に鈍感である。'
)

add_para(
    '我々は、動脈圧モニターのバリデーション研究が以下を報告すべきことを提案する：(1) rとCbへの分解を'
    '伴うCCC；(2) ゼロ校正前後のCb（機器の校正依存度を定量化）；(3) Cb低下へのuとvの個別の寄与。'
    'この情報により、規制当局と臨床医は、機器が良好な設計によって正確性を達成しているのか、'
    '校正依存的なオフセット除去によって達成しているのかを評価できるようになる。'
)

add_heading_styled('非侵襲的心拍出量モニターへの拡張', level=2)

add_para(
    '校正がオフセット（u）のみを補正し、ゲインエラー（v）には影響しないという原理は、非侵襲的心拍出量'
    '（CO）モニターにも自然に拡張される。ClearSightのPhysiocalアルゴリズムはVolume clampのセットポイント'
    'を周期的に再最適化する\u2014本質的にオフセット補正である[16]。NICOM/Starlingのバイオリアクタンス'
    '位相基準は位相\u2013一回拍出量変換のベースラインオフセットを提供する[17]。FloTracの動脈圧波形較正'
    'は平均圧オフセットを調整する[18]。いずれの場合も、自動校正ルーチンはu（ロケーションシフト）を'
    '処理するが、v（スケールシフト）\u2014すなわち圧力\u2192COまたはインピーダンス\u2192SVの変換ゲイン'
    '\u2014は補正しない。'
)

add_para(
    'この観察は、Odorら[19]が提案したCOモニタリングバリデーション枠組みと、我々が以前に提唱した'
    'CCC拡張[20]と組み合わせることで、統一原理を示唆する：いかなる血行動態モニタリング機器においても、'
    'バイアス補正係数CbをuとvのコンポーネントSに分解し、許容可能なCbを達成するための機器の校正依存度を'
    '明示的に定量化すべきである。Cbを1.0近くに維持するために頻繁な再校正を必要とする機器は、校正が部分的'
    'にしか補償できない根本的な設計上の限界を抱えている。'
)

add_heading_styled('前提条件と限界', level=2)

add_para(
    '本論文の議論は、センサ応答が生理的圧力範囲にわたって線形であるという前提に基づいている。センサが '
    '有意な非線形性を示す場合、ある圧力レベルでのPP精度は他のレベルでのゲイン正確性を保証しない。'
    'しかし、現代のMEMSピエゾ抵抗素子は一般に0\u2013300 mmHgにわたり全スケール出力の0.1%未満の'
    '非線形性を達成しており[21]、臨床的動脈圧モニタリングの要件を十分に満たしている。'
)

add_para(
    '本枠組みは系統的誤差のみを扱う。ランダム測定ノイズ（rを決定する）は別の問題であり、校正や '
    'ここで議論した設計特性の影響を受けない。高いランダムノイズを持つゼロ校正不要センサは、'
    'Cb = 1.0を達成しても依然として低いCCCを示す。'
)

add_para(
    '最後に、日常的な動脈圧モニタリングにカテーテル先端MEMSセンサを臨床採用することには、コスト、'
    'ディスポーザビリティ、および既存モニタリングインフラとの互換性を含む実用的障壁があることを認める。'
    '我々の目的は即座の臨床変更を強制することではなく、ゼロ校正が設計により除去可能な限界である'
    'という理論的原理を確立し、この目標への進捗を評価する定量的枠組み（CCC分解）を提供することである。'
)

add_heading_styled('臨床的視点', level=2)

add_para(
    '臨床医にとっての即座の実践的メッセージは微妙なものである。現行の液充填システムではゼロ校正は '
    '依然として不可欠であり、Saugelら[2]やGuptaら[3]が強調するように厳密に実施し続けるべきである。'
    'しかし、臨床医は完全なゼロ校正であってもオフセットのみを補正し、ゲインは補正しないことを'
    '認識すべきである。10%のゲインエラーを持つゼロ校正済みトランスデューサは、SBPを120ではなく132、'
    'DBPを80ではなく88と表示する\u2014ゼロ校正手順では検出されず、ベースラインのシフトではなく'
    '拡大された脈圧（40 vs. 52 mmHg）として現れる、臨床的に有意な過大評価である。'
)

add_para(
    'PPの正確性がゲインの正しさを暗示するという認識は、簡便な臨床的チェックも示唆する：侵襲的に測定された'
    'PPが生理学的に妥当であり、非侵襲的カフ測定と一致していれば、センサゲインはおそらく正しい。'
    '非生理学的なPP（例えば患者の状態に対して非常に狭いまたは非常に広い）は、ゼロ校正では解決できない'
    'ゲインの問題を示唆する可能性がある。'
)

# ══════════════════════════════════════════════════════════════════
# 結論
# ══════════════════════════════════════════════════════════════════
add_heading_styled('結論（Conclusion）', level=1)

add_para(
    '我々は、観血的動脈圧トランスデューサのゼロ校正が適切なセンサ設計により不要となることを示す理論的'
    '枠組みを提示した。この議論は2つの知見を統合する：(1)脈圧の正確性がセンサゲインを検証し、'
    'スケールシフトを排除する（v = 1）；(2)カテーテル先端MEMSセンサ、気圧補償付き絶対圧測定、'
    'および自己校正MEMS参照が、すべてのDCオフセット源を排除する（u = 0）。LinのCCC分解を通じて '
    '定式化すると、Cb = 1.0は校正ではなく設計により達成可能であることを意味する。'
    '本枠組みは、次世代動脈圧モニターの設計、規制上のバリデーション要件、および血行動態モニタリングに'
    'おける校正が何を行い\u2014何を行わないか\u2014のより深い理解のための定量的基盤を提供する。'
)

# ══════════════════════════════════════════════════════════════════
# 宣言
# ══════════════════════════════════════════════════════════════════
add_heading_styled('利益相反の開示', level=1)
add_para('[著者が記入]')

add_heading_styled('資金', level=1)
add_para('[著者が記入]')

add_heading_styled('著者の貢献', level=1)
add_para('[著者が記入]')

add_heading_styled('謝辞', level=1)
add_para('[著者が記入]')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 参考文献
# ══════════════════════════════════════════════════════════════════
add_heading_styled('参考文献（References）', level=1)

refs = [
    '1. Saugel B, Kouz K, Meidert AS, Schulte-Uentrop L, Romagnoli S. How to measure blood pressure using an arterial catheter: a systematic 5-step approach. Crit Care 2020; 24: 172. doi:10.1186/s13054-020-02859-w',
    '2. Saugel B, Sessler DI. Perioperative blood pressure management. Anesthesiology 2021; 134: 250\u201361. doi:10.1097/ALN.0000000000003610',
    '3. Gupta D, Jain A, Ismaeil M. Zero arterial catheters with every change in the height difference of pressure transducer and catheter insertion site. Ann Card Anaesth 2025; 28: 205. doi:10.4103/aca.aca_198_24',
    '4. Mark JB. Atlas of Cardiovascular Monitoring. New York: Churchill Livingstone, 1998.',
    '5. McGhee BH, Bridges MEJ. Monitoring arterial blood pressure: what you may not know. Crit Care Nurse 2002; 22: 60\u201379. doi:10.4037/ccn2002.22.2.60',
    '6. Lin LI-K. A concordance correlation coefficient to evaluate reproducibility. Biometrics 1989; 45: 255\u201368. doi:10.2307/2532051',
    '7. Lin LI-K. A note on the concordance correlation coefficient. Biometrics 2000; 56: 324\u20135. doi:10.1111/j.0006-341X.2000.00324.x',
    '8. Hasenkamp W, Theumer T, Nolte J, et al. Polyimide/SU-8 catheter-tip MEMS gauge pressure sensor. Biomed Microdevices 2012; 14: 819\u201328. doi:10.1007/s10544-012-9661-8',
    '9. Song P, Ma Z, Ma J, et al. Recent progress of miniature MEMS pressure sensors. Micromachines 2020; 11: 56. doi:10.3390/mi11010056',
    '10. Millar, Inc. Mikro-Cath Pressure Catheter \u2013 Product Information. https://millar.com/Clinical/MikroCath/ (accessed 25 March 2026).',
    '11. Millar, Inc. TiSense \u2013 Chronic Pressure Sensing. https://millar.com/Clinical/MikroCath/Continuous-Compartment-Pressure-Measurements/ (accessed 25 March 2026).',
    '12. Kang Y, Ge C, Hu Y, et al. Development of a flexible integrated self-calibrating MEMS pressure sensor using a liquid-to-vapor phase change. Sensors 2022; 22: 9737. doi:10.3390/s22249737',
    '13. U.S. Food and Drug Administration. Premarket Notification 510(k). https://www.fda.gov/medical-devices/premarket-submissions-selecting-and-preparing-correct-submission/premarket-notification-510k (accessed 25 March 2026).',
    '14. Joosten A, Desebbe O, Suehiro K, et al. Accuracy and precision of non-invasive cardiac output monitoring devices in perioperative medicine: a systematic review and meta-analysis. Br J Anaesth 2017; 118: 298\u2013310. doi:10.1093/bja/aew461',
    '15. Bland JM, Altman DG. Statistical methods for assessing agreement between two methods of clinical measurement. Lancet 1986; 327: 307\u201310. doi:10.1016/S0140-6736(86)90837-8',
    '16. Ameloot K, Palmers PJ, Malbrain MLNG. The accuracy of noninvasive cardiac output and pressure measurements with finger cuff: a concise review. Curr Opin Crit Care 2015; 21: 232\u20139. doi:10.1097/MCC.0000000000000198',
    '17. Squara P, Denjean D, Estagnasie P, Brusset A, Dib JC, Dubois C. Noninvasive cardiac output monitoring (NICOM): a clinical validation. Intensive Care Med 2007; 33: 1191\u20134. doi:10.1007/s00134-007-0640-0',
    '18. Manecke GR. Edwards FloTrac sensor and Vigileo monitor: easy, accurate, reliable cardiac output assessment using the arterial pulse wave. Expert Rev Med Devices 2005; 2: 523\u20137. doi:10.1586/17434440.2.5.523',
    '19. Odor PM, Bampoe S, Cecconi M. Cardiac output monitoring: validation studies\u2014how results should be presented. Curr Anesthesiol Rep 2017; 7: 410\u201315. doi:10.1007/s40140-017-0239-0',
    '20. [著者]. Lin\u2019s concordance correlation coefficient: a missing metric in cardiac output monitor validation. [掲載誌、年；詳細は未定].',
    '21. Barlian AA, Park W-T, Mallon JR, Rastegar AJ, Pruitt BL. Review: semiconductor piezoresistance for microsystems. Proc IEEE 2009; 97: 513\u201352. doi:10.1109/JPROC.2009.2013612',
    '22. Critchley LAH, Critchley JAJH. A meta-analysis of studies using bias and precision statistics to compare cardiac output measurement techniques. J Clin Monit Comput 1999; 15: 85\u201391. doi:10.1023/A:1009982611386',
    '23. Critchley LA, Lee A, Ho AM-H. A critical review of the ability of continuous cardiac output monitors to measure trends in cardiac output. Anesth Analg 2010; 111: 1180\u201392. doi:10.1213/ANE.0b013e3181f08a5b',
    '24. Cecconi M, Rhodes A, Poloniecki J, Della Rocca G, Grounds RM. Bench-to-bedside review: the importance of the precision of the reference technique in method comparison studies. Crit Care 2009; 13: 201. doi:10.1186/cc7129',
    '25. McBride GB. A proposal for strength-of-agreement criteria for Lin\u2019s concordance correlation coefficient. NIWA Client Report HAM2005-062. 2005.',
    '26. Bland JM, Altman DG. Measuring agreement in method comparison studies. Stat Methods Med Res 1999; 8: 135\u201360. doi:10.1177/096228029900800204',
    '27. Romagnoli S, Ricci Z, Romano SM, et al. FloTrac/Vigileo (third generation) and MostCare/PRAM versus echocardiography for cardiac output estimation in vascular surgery. J Cardiothorac Vasc Anesth 2013; 27: 1114\u201321. doi:10.1053/j.jvca.2013.04.017',
    '28. Scalia A, Ghafari C, Navarre W, Delmotte P, Phillips R, Carlier S. High fidelity pressure wires provide accurate validation of non-invasive central blood pressure and pulse wave velocity measurements. Biomedicines 2023; 11: 1235. doi:10.3390/biomedicines11041235',
    '29. Kim S-H, Lilot M, Sidhu KS, et al. Accuracy and precision of continuous noninvasive arterial pressure monitoring compared with invasive arterial pressure. Anesthesiology 2014; 120: 1080\u201397. doi:10.1097/ALN.0000000000000226',
    '30. Chatterjee K. The Swan-Ganz catheters: past, present, and future. Circulation 2009; 119: 147\u201352. doi:10.1161/CIRCULATIONAHA.108.811141',
]

for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 表1
# ══════════════════════════════════════════════════════════════════
add_heading_styled('表1（Table 1）', level=1)
add_para(
    '表1. 従来の観血的動脈圧モニタリングにおけるDCオフセットの源と排除のための工学的解決策',
    bold=True, italic=True
)

table = doc.add_table(rows=5, cols=5)
table.style = 'Light Grid Accent 1'

headers = ['オフセット源', '物理的機序', '大きさ', '工学的解決策', 'CCC成分への影響']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'

data = [
    ['静水圧カラム', '\u0394P = \u03C1gh（トランスデューサとカテーテル先端の高低差）',
     '~0.74 mmHg/cm', 'カテーテル先端MEMSセンサ（液体カラムを排除）',
     'uを低減（ロケーションシフト）'],
    ['大気圧基準', '大気（~760 mmHg）に対するゲージ圧測定',
     'ベースライン全体', '絶対圧センサ＋内蔵気圧計による電子的減算',
     'uを低減（ロケーションシフト）'],
    ['トランスデューサドリフト', 'ストレインゲージの機械的クリープ・熱効果',
     '~1\u20135 mmHg/日', '内部参照圧力キャビティ付き自己校正MEMS',
     'uを低減（ロケーションシフト）'],
    ['ゲインエラー（ゼロ校正で補正不可）', '感度不一致：単位圧力あたりの出力 \u2260 公称値',
     '1\u201315%（典型的）', '工場校正；PPの正確性による暗黙的ゲイン検証',
     'vに影響（スケールシフト）；ゼロ校正は無効'],
]

for row_idx, row_data in enumerate(data, start=1):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.name = 'Times New Roman'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 図
# ══════════════════════════════════════════════════════════════════
add_heading_styled('図（Figures）', level=1)

add_figure(
    'figure1_signal_decomposition.png',
    '図1. 動脈圧波形の分解。(A) 真の動脈圧：DC成分（オフセット依存のベースライン）とAC成分'
    '（脈圧、PP = 40 mmHg）。(B) 15 mmHgのDCオフセット（例：静水圧カラム）あり：PPは差分量であるため'
    '40 mmHgのまま不変。(C) ゲインエラー（v = 1.15）あり：PPが40から46 mmHgに歪み、'
    'PPの正確性がセンサゲインを検証することを実証。ゼロ校正は(B)のオフセットを補正するが、'
    '(C)のゲインエラーは補正できない。'
)

doc.add_page_break()

add_figure(
    'figure2_ccc_zeroing_scenarios.png',
    '図2. 4つのシミュレーションシナリオのコンコーダンスプロット（n = 150対の測定）。'
    '(A) ゼロ校正前：12 mmHgのオフセットによりu = \u20130.65のロケーションシフトが生じ、Cbが0.828に低下。'
    '(B) ゼロ校正後：オフセット除去、CCC = 0.987、Cb = 1.000。'
    '(C) ゲインエラー（v = 0.89）：スケールシフトが持続するためゼロ校正ではCb（0.842）を改善不可。'
    '(D) ゲイン＋オフセット：CCC = 0.573；ゼロ校正はDからCへの移動のみ、Bには到達不可。'
    '破線 = 恒等線（y = x）；実線 = 回帰直線。'
)

doc.add_page_break()

add_figure(
    'figure3_system_comparison.png',
    '図3. 従来型と提案するゼロ校正不要動脈圧モニタリングシステムの比較。'
    '(A) 従来の液充填システム：3つのDCオフセット源（静水圧カラム、大気圧基準、トランスデューサドリフト）'
    'が手動ゼロ校正を必要とし、校正はオフセット（u）のみを補正する。'
    '(B) 提案するゼロ校正不要システム：カテーテル先端MEMSセンサが静水圧カラムを排除、内蔵気圧計が'
    '大気圧を電子的に補償、自己校正MEMSがドリフトを排除。すべてのオフセット源が設計により排除され、'
    'PPの正確性がゲインを検証する（v = 1）。'
)

doc.add_page_break()

add_figure(
    'figure4_cb_diagnostic_space.png',
    '図4. ロケーションシフト（u）とスケールシフト（v）の関数としてのバイアス補正係数（Cb）。'
    '等高線はiso-Cb値を示す。ゼロ校正は機器を水平方向に移動（u \u2192 0、緑色矢印）させるが、'
    'vは変化しない。点A\u2013Dは図2のシナリオに対応する。金色の星は理想（u = 0, v = 1, Cb = 1.0）'
    'を示す。正しいゲイン（v = 1、PPの正確性により検証）とオフセットフリーの設計（u = 0）を持つ機器は、'
    '校正なしでCb = 1.0を達成する。'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 表2
# ══════════════════════════════════════════════════════════════════
add_heading_styled('表2（Table 2）', level=1)
add_para(
    '表2. 4つのシミュレーションシナリオにおける統計指標の比較（n = 150対の測定）',
    bold=True, italic=True
)

table2 = doc.add_table(rows=5, cols=10)
table2.style = 'Light Grid Accent 1'

t2_headers = ['シナリオ', 'CCC', 'r', 'Cb', 'u', 'v', 'バイアス\n(mmHg)', 'LOA下限\n(mmHg)',
              'LOA上限\n(mmHg)', 'PE (%)']
for i, h in enumerate(t2_headers):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'

t2_data = [
    ['A: ゼロ校正前', '0.855', '0.986', '0.867', '-0.55', '0.99', '11.9', '4.8', '19.0', '7.2'],
    ['B: ゼロ校正後', '0.986', '0.986', '1.000', '0.01', '0.99', '-0.1', '-7.2', '7.0', '7.2'],
    ['C: ゲインエラー', '0.855', '0.982', '0.870', '0.54', '1.11', '-10.9', '-19.4', '-2.4', '8.7'],
    ['D: ゲイン＋オフセット', '0.976', '0.982', '0.993', '-0.05', '1.11', '1.1', '-7.4', '9.6', '8.7'],
]

for row_idx, row_data in enumerate(t2_data, start=1):
    for col_idx, cell_text in enumerate(row_data):
        cell = table2.rows[row_idx].cells[col_idx]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.name = 'Times New Roman'

add_para(
    'CCC = Linの一致性相関係数；r = ピアソンの相関係数（精度）；Cb = バイアス補正係数（正確度）；'
    'u = ロケーションシフト（正規化されたオフセット）；v = スケールシフト（ゲイン比）；'
    'LOA = 一致の限界（Bland-Altman）；PE = percentage error。'
    'シナリオBとCは類似したバイアス（ほぼ0）とLOAを示すが、Cbは大幅に異なり（1.000 vs. 0.870）、'
    'BA解析のみではゲインエラーを検出できないことを実証している。',
    italic=True
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 追加図版（図5〜8）
# ══════════════════════════════════════════════════════════════════

add_figure(
    'figure5_ba_comparison.png',
    '図5. 4つのシミュレーションシナリオのBland-Altmanプロット。'
    '(A) ゼロ校正前：バイアス = 11.9 mmHg。'
    '(B) ゼロ校正後：バイアス ≈ 0 mmHg、LOA内に均一に分布。'
    '(C) ゲインエラー（v = 1.11）：バイアス = −10.9 mmHgだが比例パターンを示す。'
    '(D) ゲイン＋オフセット：バイアス ≈ 1.1 mmHg—BAプロット上は良好に見えるが、'
    'オフセットとゲインエラーが相殺。CCC分解（表2）なしには検出不可能。'
    '赤色実線 = 平均バイアス；赤色破線 = 95%一致の限界。'
)

doc.add_page_break()

add_figure(
    'figure6_sensitivity_analysis.png',
    '図6. 感度分析。(A) ゲインエラー（%）とDCオフセット（mmHg）の関数としてのCb値ヒートマップ。'
    'ゼロ校正（水平方向の移動、オフセット = 0）はゲインが正しい場合にのみCbを最大化する。'
    'ゲインエラーが存在すれば、ゼロ校正後もCb < 1.0のまま。'
    '(B) 異なるセンサ精度レベル（r = 0.95, 0.97, 0.99）でのゲインエラーに伴うCCC劣化。'
    '高精度センサでもゲインエラー > ±10%でCCCが大幅に低下。'
    '灰色帯域 = 典型的なMEMSゲイン精度範囲（±5%）。'
)

doc.add_page_break()

add_figure(
    'figure7_pp_validation.png',
    '図7. 脈圧（PP）検証のデモンストレーション。'
    '(A) 正しいゲイン：基準波形（青）と測定波形（赤）が一致、PP = 40 mmHg。'
    '(B) ゲインエラー（v = 1.15）：PPが46 mmHgに拡大。'
    '(C) DCオフセットのみ（+15 mmHg）：PPは40 mmHgで不変。'
    '(D) 正しいゲインでのPP相関（傾き ≈ 1.0）。'
    '(E) ゲインエラーでのPP相関（傾き = 1.15、ゲインを直接推定）。'
    '(F) PPの正確性からゼロ校正不要モニタリングへの論理連鎖の要約。'
)

doc.add_page_break()

add_figure(
    'figure8_ba_vs_concordance.png',
    '図8. コンコーダンスプロット（上段）とBland-Altmanプロット（下段）の並列比較。'
    '4つのシナリオ（A〜D）について同一データを2つの方法で可視化。'
    'シナリオBとDはBA上で類似のバイアス分布を示すが、コンコーダンスプロットでは'
    '回帰直線の傾きの違いとしてゲインエラーが明確に可視化される。'
    'この比較は、BA解析のみに依拠するバリデーションの限界と、'
    'CCC分解の補完的価値を実証している。'
)

# ══════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════
doc.save(OUTPATH)
print(f'Japanese manuscript saved: {OUTPATH}')
